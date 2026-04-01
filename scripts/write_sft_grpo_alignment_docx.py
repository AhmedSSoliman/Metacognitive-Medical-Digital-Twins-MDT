from pathlib import Path
from datetime import datetime

from docx import Document


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = out_dir / f"sft_grpo_alignment_explanation_{timestamp}.docx"

    doc = Document()
    doc.add_heading("SFT + GRPO Training and Mixed-Dataset Alignment", level=1)

    doc.add_paragraph(
        "This document explains why we train Supervised Fine-Tuning (SFT) first and "
        "Group Relative Policy Optimization (GRPO) second, what the final model contains, "
        "and how to align two datasets when one includes chain-of-thought reasoning and the other does not."
    )

    doc.add_heading("1) Why SFT first, then GRPO", level=2)
    doc.add_paragraph(
        "SFT builds the model's core capability by teaching task format, instruction-following, "
        "medical language style, and stable response patterns from demonstration data."
    )
    doc.add_paragraph(
        "GRPO then performs alignment using reward functions (for example semantic quality, "
        "safety, empathy, metacognition, and proactivity). GRPO optimizes behavioral preferences "
        "that are difficult to fully encode with supervised labels alone."
    )
    doc.add_paragraph("A practical summary:")
    doc.add_paragraph("SFT = learn the task", style="List Bullet")
    doc.add_paragraph("GRPO = optimize preferred behavior", style="List Bullet")

    doc.add_heading("2) What the final model is", level=2)
    doc.add_paragraph(
        "Yes, the final aligned model combines both stages. GRPO starts from SFT weights, "
        "so the GRPO checkpoint contains SFT-learned capability plus RL-based alignment refinements."
    )
    doc.add_paragraph(
        "In deployment, the default choice is usually the post-GRPO checkpoint (unless validation "
        "shows instability and you intentionally fall back to SFT-only)."
    )

    doc.add_heading("3) Alignment with two datasets (CoT vs non-CoT)", level=2)
    doc.add_paragraph(
        "When one dataset includes chain-of-thought (CoT/reasoning) and the other does not, "
        "treat CoT as optional supervision rather than a mandatory field for all samples."
    )
    doc.add_paragraph("Recommended approach:")
    doc.add_paragraph("Unify schema on shared fields (prompt, response, metadata).", style="List Bullet")
    doc.add_paragraph("Use optional/masked training targets for reasoning fields (e.g., think).", style="List Bullet")
    doc.add_paragraph("Do not penalize missing reasoning on non-CoT samples.", style="List Bullet")
    doc.add_paragraph("Optimize global quality rewards on final responses across both datasets.", style="List Bullet")

    doc.add_heading("4) Stage-specific guidance", level=2)
    doc.add_paragraph("During SFT:")
    doc.add_paragraph("For CoT data: learn available reasoning structure.", style="List Bullet")
    doc.add_paragraph("For non-CoT data: train final-answer behavior without forcing reasoning text.", style="List Bullet")
    doc.add_paragraph("Use loss masking for missing reasoning fields.", style="List Bullet")

    doc.add_paragraph("During GRPO:")
    doc.add_paragraph("Apply semantic/safety/empathy rewards across all samples.", style="List Bullet")
    doc.add_paragraph("Gate reasoning-format rewards only when reasoning is present/required.", style="List Bullet")
    doc.add_paragraph("Prioritize outcome quality over mandatory visible CoT to avoid over-generation of reasoning.", style="List Bullet")

    doc.add_heading("5) Key caution", level=2)
    doc.add_paragraph(
        "If reward design strongly favors explicit chain-of-thought text everywhere, the model may "
        "overproduce reasoning even when concise answers are preferred. Align rewards with deployment policy."
    )

    doc.add_heading("6) Short conclusions", level=2)
    doc.add_paragraph("Purpose of SFT -> GRPO: capability first, alignment second.", style="List Bullet")
    doc.add_paragraph("Final model: yes, one model containing both stages (GRPO-over-SFT).", style="List Bullet")
    doc.add_paragraph("Mixed data alignment: shared output objective + optional CoT supervision.", style="List Bullet")

    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
