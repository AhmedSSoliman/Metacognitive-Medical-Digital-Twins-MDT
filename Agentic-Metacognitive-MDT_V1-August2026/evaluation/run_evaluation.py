"""
evaluation/run_evaluation.py

Checkpoint evaluation driver: loads a trained model, generates over an
evaluation prompt set, scores every endpoint via evaluation/report.py, and
renders a PNG/DOCX visual report.

PORTING NOTE (2026-08-12): ported from the source repo's
scripts/run_evaluation.py (../Agentic-DT_V1-July/scripts/run_evaluation.py).

PLACEMENT JUDGMENT CALL: the source kept this under scripts/. In this layout
scripts/ holds only shell/SBATCH wrappers, and this is a substantive Python
module (~330 lines) that tests import from directly -- so it moved into
evaluation/ alongside the endpoints it drives. scripts/run_evaluation.sh now
invokes it as `python -m evaluation.run_evaluation`.

Imports rewritten to core.parsing / core.hypergraph.verification /
evaluation.report, and the model import deferred into main() (see the comment
at the import site).

ORIGINAL scripts/run_evaluation.py MODULE DOCSTRING:

    scripts/run_evaluation.py

    End-to-end evaluation driver: loads a trained checkpoint, runs it against a
    held-out EVALUATION partition (never the derivation partition used to mine
    the hypergraph), extracts everything evaluation/metrics.py's
    run_full_evaluation() needs, and saves a report.

    This did not exist as a runnable script before -- evaluation/metrics.py
    defined all six endpoint functions and run_full_evaluation() to combine
    them, but nothing actually called that function against a real checkpoint's
    output. This script is the missing "glue": it is the difference between
    "the evaluation logic exists" and "you can actually run an evaluation."

    WHAT THIS SCRIPT DOES:
      1. Loads the checkpoint and the evaluation-partition prompt dataset
         (built the same way as Phase 2's training prompts, via
         data/build_grpo_prompts.py, but run against the EVALUATION cohort/
         vitals/labs -- never the derivation ones the hypergraph was mined from).
      2. Generates one completion per example, timing each generation for the
         Deployment Feasibility endpoint's latency metrics.
      3. Extracts a `risk_score` (continuous) for the Deterioration Detection
         endpoint via a simple, documented heuristic (see `extract_risk_score`)
         since neither this project's training pipeline nor its prior repo
         currently trains a proper regression head end-to-end (ForecastHead in
         models/multi_stream.py is defined but not yet wired into any training
         script -- see that class's docstring and the project README for the
         open question of whether to train it or rely on this heuristic
         instead).
      4. Calls evaluation.report.run_full_evaluation() and saves the report.
"""


from __future__ import annotations

import argparse
import json
import logging
import re
import time
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")  # headless backend -- must precede pyplot import; this runs inside a
                        # SLURM batch job with no $DISPLAY, and forgetting this can hang/crash
                        # on cluster nodes
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from sklearn.metrics import confusion_matrix, roc_curve

from core.parsing import parse_streams
from core.hypergraph.verification import InterimRuleBasedChecker, LearnedHypergraphChecker
from evaluation.report import EvaluationResult, run_full_evaluation, save_evaluation_report

# MultiStreamModel/MultiStreamConfig import torch+unsloth transitively, so they
# are imported LAZILY inside main() rather than at module level. This is a
# deliberate change from the source script: tests/test_evaluation/
# test_evaluation_metrics.py's Section E imports generate_visual_report from
# this module to test the matplotlib/docx report rendering, which needs no
# model at all -- and under the source layout it could only do that by
# side-loading the file via importlib. Deferring the model import lets this
# module be imported normally without a GPU, while `main()` behaves identically.

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# risk_score extraction heuristic
#
# This is an HONEST, documented placeholder, not a claim that this is the
# "right" way to get a risk score -- it is a simple keyword-count severity
# proxy over the <patient_state> text, used because nothing in this
# project's current training pipeline produces a proper calibrated
# probability. If you train the ForecastHead regression module (or extend
# GRPO training to jointly optimize it), replace this function's body with
# a call into that head's output instead -- the rest of this script
# (everything downstream of extract_risk_score) does not need to change.
# ---------------------------------------------------------------------------

_SEVERITY_TERMS = {
    "shock": 1.0, "septic": 0.9, "critical": 0.9, "unstable": 0.8,
    "deteriorat": 0.8, "hypotens": 0.6, "hyperlactatemi": 0.6,
    "tachycardi": 0.4, "concerning": 0.5, "urgent": 0.7,
    "stable": -0.5, "improving": -0.6, "reassuring": -0.6, "normal": -0.3,
}


def extract_risk_score(patient_state_text: str) -> float:
    """Crude keyword-weighted severity heuristic, min-max clamped to [0, 1].
    See the module-level comment above for why this exists and what to
    replace it with once a trained regression head is available.
    """
    if not patient_state_text:
        return 0.5  # neutral default when there's nothing to score at all
    text = patient_state_text.lower()
    score = 0.5  # start neutral, push up/down based on matched terms
    for term, weight in _SEVERITY_TERMS.items():
        if re.search(term, text):
            score += weight * 0.15  # each matched term nudges the score, doesn't fully determine it
    return float(np.clip(score, 0.0, 1.0))


# ---------------------------------------------------------------------------
# Visual report generation (PNG + docx)
#
# run_full_evaluation() only returns a summary EvaluationResult, not the raw
# risk_score/true_deterioration_label arrays needed to rebuild a confusion
# matrix or ROC curve -- so this function needs eval_df passed in separately,
# the same eval_df used to produce `result`. `threshold` MUST match the
# deterioration_threshold passed to run_full_evaluation(), or the confusion
# matrix panel will disagree with the reported precision/recall/F2 numbers;
# main() below uses a single args.deterioration_threshold as the source of
# truth for both, so this can't drift in normal use.
# ---------------------------------------------------------------------------

def generate_visual_report(result: EvaluationResult, eval_df: pd.DataFrame, output_dir: Path,
                            report_name: str = "evaluation_report", threshold: float = 0.5) -> tuple[Path, Path]:
    """Builds a 3-panel matplotlib PNG (confusion matrix, ROC curve, endpoint-
    score bar chart) and a docx summary report. Returns (png_path, docx_path).
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    png_path = output_dir / f"{report_name}.png"
    docx_path = output_dir / f"{report_name}.docx"

    y_true = eval_df["true_deterioration_label"].to_numpy()
    y_score = eval_df["risk_score"].to_numpy()
    y_pred = (y_score >= threshold).astype(int)

    fig, axes = plt.subplots(1, 3, figsize=(18, 5))

    # Panel 1: confusion matrix. Plain imshow + text annotations rather than
    # seaborn -- a single 2x2 heatmap doesn't need an extra dependency.
    cm = confusion_matrix(y_true, y_pred, labels=[0, 1])
    axes[0].imshow(cm, cmap="Blues")
    axes[0].set_title(f"Confusion Matrix (threshold={threshold})")
    axes[0].set_xticks([0, 1]); axes[0].set_xticklabels(["No Deterioration", "Deterioration"])
    axes[0].set_yticks([0, 1]); axes[0].set_yticklabels(["No Deterioration", "Deterioration"])
    axes[0].set_xlabel("Predicted"); axes[0].set_ylabel("Actual")
    cm_max = cm.max() if cm.size else 0
    for i in range(cm.shape[0]):
        for j in range(cm.shape[1]):
            color = "white" if cm_max and cm[i, j] > cm_max / 2 else "black"
            axes[0].text(j, i, str(cm[i, j]), ha="center", va="center", color=color, fontsize=14)

    # Panel 2: ROC curve, annotated with the bootstrapped AUC CI so it's
    # visible in the report, not just buried in the JSON.
    det = result.deterioration_detection
    if len(np.unique(y_true)) > 1:
        fpr, tpr, _ = roc_curve(y_true, y_score)
        axes[1].plot(fpr, tpr, color="darkorange", lw=2,
                     label=f"AUROC = {result.deterioration_auroc:.3f} "
                           f"[{det['auroc_ci_lower']:.3f}, {det['auroc_ci_upper']:.3f}]")
        axes[1].plot([0, 1], [0, 1], color="navy", lw=1, linestyle="--")
        axes[1].legend(loc="lower right")
    axes[1].set_title("ROC Curve")
    axes[1].set_xlabel("False Positive Rate"); axes[1].set_ylabel("True Positive Rate")

    # Panel 3: bar chart of the endpoint scores that actually exist in
    # EvaluationResult today -- deliberately NOT the reference codebase's
    # keyword-based scores. trajectory_mae is excluded from the chart (it's
    # an unbounded, lower-is-better error metric, not a [0,1] score, so
    # plotting it on the same axis as the others would be misleading) and
    # shown as a text annotation instead. Endpoint 4 (HITL Audit Latency)
    # has no evaluate_* function or EvaluationResult field yet, so it can't
    # be plotted -- this does not invent one.
    metric_labels = ["Topological\nFidelity", "Deterioration\nAUROC", "Structural\nEmpathy",
                      "Format\nCompliance", "Schema\nValidity"]
    metric_values = [result.topological_fidelity, result.deterioration_auroc,
                      result.structural_empathy_mean, result.format_compliance_rate,
                      result.deployment["schema_validity_rate"]]
    bars = axes[2].bar(metric_labels, metric_values,
                        color=["#4C72B0", "#55A868", "#8172B2", "#CCB974", "#64B5CD"])
    axes[2].set_title("Core Performance Metrics")
    axes[2].set_ylim(0, 1.15)
    for bar in bars:
        height = bar.get_height()
        axes[2].annotate(f"{height:.2f}", xy=(bar.get_x() + bar.get_width() / 2, height),
                          xytext=(0, 3), textcoords="offset points", ha="center", va="bottom")
    mae_text = f"{result.trajectory_mae:.3f}" if result.trajectory_mae is not None else "not computed (no forecast columns present)"
    fig.text(0.5, -0.03, f"Trajectory MAE: {mae_text}", ha="center", fontsize=10)

    plt.tight_layout()
    plt.savefig(png_path, dpi=200, bbox_inches="tight")
    plt.close(fig)

    # docx report: title, a metrics table covering every EvaluationResult
    # field (flattening the nested deployment/deterioration_detection dicts
    # into dotted row labels), and the panel image embedded below it.
    doc = Document()
    doc.add_heading("Evaluation Report", level=0)
    doc.add_paragraph(f"n_examples: {result.n_examples}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"

    def _fmt(v):
        if v is None or (isinstance(v, float) and np.isnan(v)):
            return "N/A"
        if isinstance(v, float):
            return f"{v:.4f}"
        return str(v)

    for key, value in result.__dict__.items():
        if isinstance(value, dict):
            for sub_key, sub_value in value.items():
                row = table.add_row().cells
                row[0].text = f"{key}.{sub_key}"
                row[1].text = _fmt(sub_value)
        else:
            row = table.add_row().cells
            row[0].text = key
            row[1].text = _fmt(value)

    doc.add_paragraph("")
    doc.add_heading("Visualizations", level=1)
    doc.add_picture(str(png_path), width=Inches(6.5))

    doc.save(docx_path)
    return png_path, docx_path


def main():
    parser = argparse.ArgumentParser(description="Run the full six-endpoint evaluation against a checkpoint.")
    parser.add_argument("--checkpoint", required=True, help="Path to the model checkpoint to evaluate.")
    parser.add_argument("--eval_prompt_dataset", required=True,
                        help="Output of core/cohort/grpo_prompts.py, run against the EVALUATION "
                             "(not derivation) partition.")
    parser.add_argument("--hypergraph_mode", choices=["interim", "learned"], default="interim")
    parser.add_argument("--hypergraph_path", default="")
    parser.add_argument("--output_report", default="./evaluation/reports/evaluation_report.json")
    parser.add_argument("--max_examples", type=int, default=None,
                        help="Optionally cap the number of examples evaluated, for a quick smoke "
                             "run before committing to a full (potentially slow) evaluation pass.")
    parser.add_argument("--max_new_tokens", type=int, default=768)
    parser.add_argument("--deterioration_threshold", type=float, default=0.5,
                        help="Risk-score cutoff for classifying a prediction as positive when "
                             "computing precision/recall/F2 for the Deterioration Detection "
                             "endpoint. AUROC/AUPRC are threshold-free and unaffected by this.")
    parser.add_argument("--n_bootstrap_ci", type=int, default=500,
                        help="Number of bootstrap resamples for the AUROC confidence interval.")
    parser.add_argument("--ci_alpha", type=float, default=0.95,
                        help="Confidence level for the bootstrapped AUROC interval.")
    parser.add_argument("--skip_visual_report", action="store_true",
                        help="Skip PNG/docx report generation (the JSON report is always saved). "
                             "Useful for quick --max_examples smoke runs.")
    args = parser.parse_args()

    if args.hypergraph_mode == "interim":
        hypergraph_checker = InterimRuleBasedChecker()
    else:
        if not args.hypergraph_path:
            raise ValueError("--hypergraph_path required when --hypergraph_mode learned")
        hypergraph_checker = LearnedHypergraphChecker(args.hypergraph_path)

    logger.info("Loading checkpoint: %s", args.checkpoint)
    from training.backbone import MultiStreamModel, MultiStreamConfig
    ts_model = MultiStreamModel(MultiStreamConfig(base_model_name=args.checkpoint))

    df = pd.read_parquet(args.eval_prompt_dataset)
    if args.max_examples is not None:
        df = df.head(args.max_examples)
    logger.info("Evaluating against %d examples from %s", len(df), args.eval_prompt_dataset)

    records = []
    for i, row in df.iterrows():
        start = time.time()
        # num_return_sequences=1: evaluation runs a single generation per
        # example, not a sampled group -- group sampling (as GRPO training
        # uses) is for computing a relative advantage during training, which
        # has no meaning here; evaluation cares about what the model
        # actually outputs when asked once, the way a real deployment would.
        generations = ts_model.generate(row["prompt"], max_new_tokens=args.max_new_tokens,
                                          temperature=0.0, num_return_sequences=1)
        latency_ms = (time.time() - start) * 1000
        raw_generation = generations[0]
        parsed = parse_streams(raw_generation)

        records.append({
            "generation": raw_generation,
            "true_deterioration_label": row["deterioration_6h"],
            "risk_score": extract_risk_score(parsed.patient_state),
            "recipient_type": row["recipient_type"],
            "latency_ms": latency_ms,
            "schema_valid": parsed.well_formed,
            # Tool-call success isn't meaningfully assessable outside the
            # agentic Phase 4 rollout context (see agents/tool_use.py) --
            # defaults to True (no tool call attempted = trivially "didn't
            # fail") rather than being scored here; run evaluation through
            # agents/rollout_service.py's tool-agent path instead if you
            # need real tool-call success rates as part of this report.
            "tool_success": True,
        })

        if (i + 1) % 20 == 0:
            logger.info("Evaluated %d/%d examples ...", i + 1, len(df))

    eval_df = pd.DataFrame(records)
    result = run_full_evaluation(eval_df, hypergraph_checker,
                                  deterioration_threshold=args.deterioration_threshold,
                                  n_bootstrap=args.n_bootstrap_ci, ci_alpha=args.ci_alpha)

    output_report_path = Path(args.output_report)
    output_report_path.parent.mkdir(parents=True, exist_ok=True)
    save_evaluation_report(result, args.output_report)

    det = result.deterioration_detection
    print("\n=== Evaluation Summary ===")
    print(f"n_examples:              {result.n_examples}")
    print(f"Topological Fidelity:    {result.topological_fidelity:.3f}")
    print(f"Deterioration AUROC:     {result.deterioration_auroc:.3f} "
          f"[{det['auroc_ci_lower']:.3f}, {det['auroc_ci_upper']:.3f}] (95% bootstrap CI)")
    print(f"Deterioration AUPRC:     {result.deterioration_auprc:.3f}")
    print(f"Deterioration Precision: {det['precision']:.3f}  Recall: {det['recall']:.3f}  "
          f"F2: {det['f2_score']:.3f}  (threshold={det['threshold']})")
    print(f"Trajectory MAE:          {result.trajectory_mae}")
    print(f"Structural Empathy:      {result.structural_empathy_mean:.3f}")
    print(f"Format Compliance Rate:  {result.format_compliance_rate:.1%}")
    print(f"Deployment metrics:      {result.deployment}")
    print(f"\nFull report saved to: {args.output_report}")

    if not args.skip_visual_report:
        png_path, docx_path = generate_visual_report(
            result, eval_df, output_dir=output_report_path.parent,
            report_name=output_report_path.stem, threshold=args.deterioration_threshold,
        )
        print(f"Visual report saved to: {png_path}, {docx_path}")


if __name__ == "__main__":
    main()
